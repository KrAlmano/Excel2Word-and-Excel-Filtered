import os
import shutil
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import pandas as pd
import openpyxl
import re
import numpy as np
import sys
from docx import Document
import customtkinter as ctk

# Veritabanı dizinini masaüstünde oluşturur
DATABASE_DIR = os.path.join(os.path.expanduser('~/Desktop'), 'ExcelVeritabaniV2')

class ExcelYoneticiV2(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Excel Yöneticisi V2")
        self.geometry("1300x800")
        ctk.set_appearance_mode("System")  # Sistem temasını kullan (Dark/Light)
        ctk.set_default_color_theme("blue")  # Tema rengi

        # Sınıf değişkenleri
        self.df = None
        self.df_filtered = None
        self.secili_dosya = None
        self.secili_klasor = None
        self.filtreler = []  # [(kolon, operator, deger1, deger2, tip)]

        # Arayüzü oluştur
        self._arayuz_olustur()
        self._dosya_agacini_guncelle()

        # Sağ tık menüsü
        self.sag_menu = tk.Menu(self, tearoff=0)
        self.sag_menu.add_command(label="Sil", command=self._dosya_sil)
        self.sag_menu.add_command(label="Yeniden Adlandır", command=self._dosya_yeniden_adlandir)

    def _arayuz_olustur(self):
        # Ana çerçeve
        ana_frame = ctk.CTkFrame(self)
        ana_frame.pack(fill=ctk.BOTH, expand=True, padx=10, pady=10)

        # Sol panel: Dosya ağacı
        sol_frame = ctk.CTkFrame(ana_frame, width=250)
        sol_frame.pack(side=ctk.LEFT, fill=ctk.Y, padx=(0, 10))
        sol_frame.pack_propagate(False) # Genişliğin sabit kalmasını sağlar

        ctk.CTkLabel(sol_frame, text="Veritabanı", font=("Arial", 14, "bold")).pack(pady=10)
        
        self.dosya_tree = ttk.Treeview(sol_frame, show="tree")
        self.dosya_tree.pack(fill=ctk.BOTH, expand=True, padx=5)
        self.dosya_tree.bind('<<TreeviewSelect>>', self._dosya_agacinda_secildi)
        self.dosya_tree.bind('<Button-3>', self._treeview_sag_tik_menu)
        
        ctk.CTkButton(sol_frame, text="Excel Yükle", command=self.dosya_yukle).pack(pady=10, padx=5, fill='x')

        # Orta panel: Filtre ve tablo
        orta_frame = ctk.CTkFrame(ana_frame)
        orta_frame.pack(side=ctk.LEFT, fill=ctk.BOTH, expand=True)
        
        self._olustur_filtre_paneli(orta_frame)
        
        self.tablo_frame = ctk.CTkFrame(orta_frame)
        self.tablo_frame.pack(fill=ctk.BOTH, expand=True, padx=5, pady=5)
        
        self.tree = ttk.Treeview(self.tablo_frame, show="headings")
        self.vsb = ttk.Scrollbar(self.tablo_frame, orient="vertical", command=self.tree.yview)
        self.hsb = ttk.Scrollbar(self.tablo_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=self.vsb.set, xscrollcommand=self.hsb.set)

        self.tree.grid(row=0, column=0, sticky='nsew')
        self.vsb.grid(row=0, column=1, sticky='ns')
        self.hsb.grid(row=1, column=0, sticky='ew')
        
        self.tablo_frame.grid_rowconfigure(0, weight=1)
        self.tablo_frame.grid_columnconfigure(0, weight=1)

    def _olustur_filtre_paneli(self, parent):
        self.filtre_frame = ctk.CTkFrame(parent)
        self.filtre_frame.pack(fill=ctk.X, pady=5, padx=5)
        
        ctk.CTkLabel(self.filtre_frame, text="Filtreleme Paneli", font=("Arial", 12, "bold")).grid(row=0, column=0, columnspan=4, sticky="w", padx=5, pady=5)
        
        # --- DÜZELTME 1: ComboBox'lara `command` parametresi eklendi ---
        self.filtre_sutun = ctk.CTkComboBox(self.filtre_frame, state="disabled", width=180, command=self._filtre_sutun_degisti)
        self.filtre_sutun.grid(row=1, column=0, padx=5, pady=5)
        
        self.filtre_operator = ctk.CTkComboBox(self.filtre_frame, state="disabled", width=140, command=self._filtre_operator_degisti)
        self.filtre_operator.grid(row=1, column=1, padx=5, pady=5)
        
        self.filtre_giris1 = ctk.CTkEntry(self.filtre_frame, width=120)
        self.filtre_giris1.grid(row=1, column=2, padx=5, pady=5)
        
        self.filtre_giris2 = ctk.CTkEntry(self.filtre_frame, width=120)
        self.filtre_giris2.grid(row=1, column=3, padx=5, pady=5)
        self.filtre_giris2.grid_remove() # Başlangıçta gizle

        self.filtre_btn = ctk.CTkButton(self.filtre_frame, text="Filtre Ekle", command=self.filtre_ekle, state=ctk.DISABLED)
        self.filtre_btn.grid(row=2, column=0, padx=5, pady=5, sticky="ew")
        
        self.filtre_temizle_btn = ctk.CTkButton(self.filtre_frame, text="Tümünü Temizle", command=self.filtre_temizle, state=ctk.DISABLED)
        self.filtre_temizle_btn.grid(row=2, column=1, padx=5, pady=5, sticky="ew")
        
        self.filtre_kaydet_btn = ctk.CTkButton(self.filtre_frame, text="Filtreli Kaydet", command=self.filtreli_veriyi_kaydet)
        self.filtre_kaydet_btn.grid(row=2, column=2, padx=5, pady=5, sticky="ew")
        
        self.veri_duzenle_btn = ctk.CTkButton(self.filtre_frame, text="Veri Düzenle", command=self.veri_duzenle_pencere)
        self.veri_duzenle_btn.grid(row=2, column=3, padx=5, pady=5, sticky="ew")
        
        self.filtre_listbox = tk.Listbox(self.filtre_frame, height=3)
        self.filtre_listbox.grid(row=3, column=0, columnspan=4, padx=5, pady=5, sticky="ew")
        
        self.filtre_sil_btn = ctk.CTkButton(self.filtre_frame, text="Seçili Filtreyi Kaldır", command=self.filtre_sil, state=ctk.DISABLED)
        self.filtre_sil_btn.grid(row=4, column=0, padx=5, pady=5, sticky="w")
        
        self.filtre_sonuc_label = ctk.CTkLabel(self.filtre_frame, text="")
        self.filtre_sonuc_label.grid(row=4, column=1, columnspan=3, sticky="w", padx=5)

    def _filtre_paneli_guncelle(self):
        if self.df is not None:
            self.filtre_sutun.configure(state="readonly", values=list(self.df.columns))
            self.filtre_sutun.set("")
            self.filtre_operator.configure(state="disabled", values=[])
            self.filtre_operator.set("")
            self.filtre_giris1.delete(0, tk.END)
            self.filtre_giris2.delete(0, tk.END)
            self.filtre_btn.configure(state="normal")
            self.filtre_temizle_btn.configure(state="normal")
            self.filtre_sil_btn.configure(state="normal")
            # --- DÜZELTME 2: .bind() satırı kaldırıldı, çünkü command kullanılıyor ---
        else:
            self.filtre_sutun.configure(state="disabled", values=[])
            self.filtre_operator.configure(state="disabled", values=[])
            self.filtre_btn.configure(state="disabled")
            self.filtre_temizle_btn.configure(state="disabled")
            self.filtre_sil_btn.configure(state="disabled")

    # --- DÜZELTME 3: Fonksiyon, olay (event) yerine seçilen değeri (secilen_sutun) alacak şekilde güncellendi ---
    def _filtre_sutun_degisti(self, secilen_sutun):
        if self.df is None or secilen_sutun not in self.df.columns:
            return
        
        self.filtre_operator.configure(state="readonly", values=['Aralık', 'Anahtar Kelime'])
        self.filtre_operator.set("")
        self.filtre_giris1.delete(0, tk.END)
        self.filtre_giris2.delete(0, tk.END)
        self.filtre_giris2.grid_remove()
        # .bind() kaldırıldı, command zaten ayarlandı.

    # --- DÜZELTME 4: Bu fonksiyon da seçilen değeri alacak şekilde güncellendi ---
    def _filtre_operator_degisti(self, secilen_operator):
        if secilen_operator == 'Aralık':
            # Orijinal grid pozisyonuna geri yerleştir
            self.filtre_giris2.grid(row=1, column=3, padx=5, pady=5)
        else:
            self.filtre_giris2.grid_remove()

    def _dosya_agacini_guncelle(self):
        self.dosya_tree.delete(*self.dosya_tree.get_children())
        if not os.path.exists(DATABASE_DIR):
            os.makedirs(DATABASE_DIR)
        for klasor in sorted(os.listdir(DATABASE_DIR)):
            klasor_yol = os.path.join(DATABASE_DIR, klasor)
            if os.path.isdir(klasor_yol):
                kid = self.dosya_tree.insert('', 'end', text=f"📁 {klasor}", open=True)
                for dosya in sorted(os.listdir(klasor_yol)):
                    if dosya.endswith('.xlsx'):
                        self.dosya_tree.insert(kid, 'end', text=f"📄 {dosya}", values=(os.path.join(klasor, dosya),))

    def dosya_yukle(self):
        dosya = filedialog.askopenfilename(title="Excel Dosyası Seç", filetypes=[("Excel Dosyaları", "*.xlsx")])
        if not dosya:
            return
        dosya_adi = os.path.basename(dosya)
        ana_isim = os.path.splitext(dosya_adi)[0]
        hedef_klasor = os.path.join(DATABASE_DIR, ana_isim)
        if not os.path.exists(hedef_klasor):
            os.makedirs(hedef_klasor)
        hedef_yol = os.path.join(hedef_klasor, dosya_adi)
        try:
            shutil.copy2(dosya, hedef_yol)
            messagebox.showinfo("Başarılı", f"{dosya_adi} başarıyla yüklendi.")
            self._dosya_agacini_guncelle()
        except Exception as e:
            messagebox.showerror("Hata", f"Dosya yüklenemedi:\n{e}")

    def _dosya_agacinda_secildi(self, event):
        secim = self.dosya_tree.selection()
        if not secim:
            return
        item = self.dosya_tree.item(secim[0])
        parent = self.dosya_tree.parent(secim[0])
        if parent:  # Dosya seçildi
            klasor = self.dosya_tree.item(parent)['text'].replace("📁 ", "")
            dosya_adi = item['text'].replace("📄 ", "")
            dosya_yolu = os.path.join(DATABASE_DIR, klasor, dosya_adi)
            self.secili_klasor = klasor
            self.secili_dosya = dosya_adi
            try:
                df = pd.read_excel(dosya_yolu, engine="openpyxl")
                df = df.dropna(how='all')
                if any([str(col).startswith('Unnamed') for col in df.columns]):
                    new_header = df.iloc[0]
                    df = df[1:]
                    df.columns = new_header
                    df = df.reset_index(drop=True)
                self.df = df
                self.df_filtered = None
                self.filtreler = []
                self._tabloyu_goster(df)
                self._filtre_paneli_guncelle()
            except Exception as e:
                messagebox.showerror("Hata", f"Dosya okunamadı:\n{e}")
        else:
            self.secili_klasor = item['text'].replace("📁 ", "")
            self.secili_dosya = None

    def _treeview_sag_tik_menu(self, event):
        iid = self.dosya_tree.identify_row(event.y)
        if iid:
            self.dosya_tree.selection_set(iid)
            try:
                self.sag_menu.tk_popup(event.x_root, event.y_root)
            finally:
                self.sag_menu.grab_release()

    def _dosya_sil(self):
        secim = self.dosya_tree.selection()
        if not secim or len(secim) != 1:
            messagebox.showwarning("Uyarı!", "Lütfen silmek için bir dosya veya klasör seçin.")
            return
        item = self.dosya_tree.item(secim[0])
        parent = self.dosya_tree.parent(secim[0])
        if parent:  # Dosya sil
            klasor = self.dosya_tree.item(parent)['text'].replace("📁 ", "")
            dosya_adi = item['text'].replace("📄 ", "")
            dosya_yolu = os.path.join(DATABASE_DIR, klasor, dosya_adi)
            if messagebox.askyesno("Sil", f"{dosya_adi} dosyasını silmek istediğinize emin misiniz?"):
                try:
                    os.remove(dosya_yolu)
                    self._dosya_agacini_guncelle()
                except Exception as e:
                    messagebox.showerror("Hata", f"Dosya silinemedi:\n{e}")
        else:  # Klasör sil
            klasor = item['text'].replace("📁 ", "")
            klasor_yolu = os.path.join(DATABASE_DIR, klasor)
            if messagebox.askyesno("Sil", f"{klasor} klasörünü ve içindeki tüm dosyaları silmek istediğinize emin misiniz?"):
                try:
                    shutil.rmtree(klasor_yolu)
                    self._dosya_agacini_guncelle()
                except Exception as e:
                    messagebox.showerror("Hata", f"Klasör silinemedi:\n{e}")

    def _dosya_yeniden_adlandir(self):
        secim = self.dosya_tree.selection()
        if not secim or len(secim) != 1:
            messagebox.showwarning("Uyarı", "Lütfen yeniden adlandırmak için bir dosya veya klasör seçin.")
            return
        item = self.dosya_tree.item(secim[0])
        parent = self.dosya_tree.parent(secim[0])
        if parent:  # Dosya yeniden adlandır
            klasor = self.dosya_tree.item(parent)['text'].replace("📁 ", "")
            eski_ad = item['text'].replace("📄 ", "")
            yeni_ad = simpledialog.askstring("Yeniden Adlandır", "Yeni dosya adını girin (uzantı dahil):", initialvalue=eski_ad)
            if yeni_ad and yeni_ad != eski_ad:
                eski_yol = os.path.join(DATABASE_DIR, klasor, eski_ad)
                yeni_yol = os.path.join(DATABASE_DIR, klasor, yeni_ad)
                try:
                    os.rename(eski_yol, yeni_yol)
                    self._dosya_agacini_guncelle()
                except Exception as e:
                    messagebox.showerror("Hata", f"Dosya yeniden adlandırılamadı:\n{e}")
        else:  # Klasör yeniden adlandır
            eski_klasor = item['text'].replace("📁 ", "")
            yeni_klasor = simpledialog.askstring("Yeniden Adlandır", "Yeni klasör adını girin:", initialvalue=eski_klasor)
            if yeni_klasor and yeni_klasor != eski_klasor:
                eski_yol = os.path.join(DATABASE_DIR, eski_klasor)
                yeni_yol = os.path.join(DATABASE_DIR, yeni_klasor)
                try:
                    os.rename(eski_yol, yeni_yol)
                    self._dosya_agacini_guncelle()
                except Exception as e:
                    messagebox.showerror("Hata", f"Klasör yeniden adlandırılamadı:\n{e}")

    def _kolon_tipi_bul(self, col):
        if self.df is None or col not in self.df.columns: return 'metin'
        seri = self.df[col].dropna().astype(str)
        if seri.empty: return 'metin'
        try:
            pd.to_datetime(seri.iloc[:10], dayfirst=True, errors='raise')
            return 'tarih'
        except (ValueError, TypeError): pass
        try:
            pd.to_numeric(seri.str.replace(r'[^0-9,.-]', '', regex=True).str.replace('.', '', regex=False).str.replace(',', '.', regex=False), errors='raise')
            return 'sayi'
        except (ValueError, TypeError): pass
        return 'metin'

    def filtre_ekle(self):
        col = self.filtre_sutun.get()
        op = self.filtre_operator.get()
        val1 = self.filtre_giris1.get().strip()
        val2 = self.filtre_giris2.get().strip() if op == 'Aralık' else None
        if not col or not op or not val1 or (op == 'Aralık' and not val2):
            messagebox.showwarning("Uyarı", "Lütfen tüm filtre alanlarını doldurun.")
            return
        tip = self._kolon_tipi_bul(col)
        self.filtreler.append((col, op, val1, val2, tip))
        self._filtre_listbox_guncelle()
        self._filtreleri_uygula()

    def _filtre_listbox_guncelle(self):
        self.filtre_listbox.delete(0, tk.END)
        for f in self.filtreler:
            col, op, val1, val2, tip = f
            if op == 'Aralık':
                self.filtre_listbox.insert(tk.END, f"{col} ({op}): {val1} - {val2}")
            else:
                self.filtre_listbox.insert(tk.END, f"{col} ({op}): {val1}")

    def filtre_sil(self):
        sec = self.filtre_listbox.curselection()
        if not sec: return
        self.filtreler.pop(sec[0])
        self._filtre_listbox_guncelle()
        self._filtreleri_uygula()

    def filtre_temizle(self):
        self.filtreler = []
        self._filtre_listbox_guncelle()
        self._filtreleri_uygula()

    def _filtreleri_uygula(self):
        if self.df is None: return
        if not self.filtreler:
            self.df_filtered = None
            self._tabloyu_goster(self.df)
            self.filtre_sonuc_label.configure(text=f"Toplam satır: {len(self.df)}")
            return
        
        df = self.df.copy()
        for col, op, val1, val2, tip in self.filtreler:
            try:
                if op == 'Aralık':
                    if tip == 'tarih':
                        df[col] = pd.to_datetime(df[col], dayfirst=True, errors='coerce')
                        v1 = pd.to_datetime(val1, dayfirst=True, errors='coerce')
                        v2 = pd.to_datetime(val2, dayfirst=True, errors='coerce')
                        if pd.notna(v1) and pd.notna(v2):
                            df = df[df[col].between(v1, v2)]
                    elif tip == 'sayi':
                        temp_col = df[col].astype(str).str.replace(r'[^\d,.-]', '', regex=True).str.replace('.', '', regex=False).str.replace(',', '.', regex=False)
                        df[col] = pd.to_numeric(temp_col, errors='coerce')
                        v1 = float(val1.replace('.', '').replace(',', '.'))
                        v2 = float(val2.replace('.', '').replace(',', '.'))
                        df = df[df[col].between(v1, v2)]
                elif op == 'Anahtar Kelime':
                    keywords = re.split(r',| ve | ya da | veya ', val1, flags=re.IGNORECASE)
                    keywords = [k.strip().lower() for k in keywords if k.strip()]
                    mask = df[col].astype(str).str.lower().apply(lambda x: all(kw in x for kw in keywords))
                    df = df[mask]
            except Exception as e:
                messagebox.showerror("Filtreleme Hatası", f"'{col}' filtresi uygulanamadı:\n{e}")
        
        self.df_filtered = df
        self._tabloyu_goster(df)
        self.filtre_sonuc_label.configure(text=f"Filtreli: {len(df)} / Toplam: {len(self.df)}")

    def filtreli_veriyi_kaydet(self):
        df_to_save = self.df_filtered if self.df_filtered is not None else self.df
        if df_to_save is None or self.secili_klasor is None or self.secili_dosya is None:
            messagebox.showwarning("Uyarı", "Kaydedilecek veri yok.")
            return
        
        ana_isim = self.secili_klasor
        orijinal_ad = os.path.splitext(self.secili_dosya)[0]
        default_name = f"{orijinal_ad}_filtreli.docx"
        yeni_ad = simpledialog.askstring("Filtreli Veriyi Kaydet", "Kaydedilecek dosya adını girin:", initialvalue=default_name)
        if not yeni_ad: return
        if not yeni_ad.endswith('.docx'): yeni_ad += '.docx'
        
        klasor_yolu = os.path.join(DATABASE_DIR, ana_isim)
        kayit_yolu = os.path.join(klasor_yolu, yeni_ad)
        
        filtre_aciklamalari = [f"Kaynak Dosya: {self.secili_dosya}"]
        if self.filtreler:
            for f in self.filtreler:
                col, op, val1, val2, _ = f
                if op == 'Aralık':
                    filtre_aciklamalari.append(f"- Filtre: {col} ({op}): {val1} - {val2}")
                else:
                    filtre_aciklamalari.append(f"- Filtre: {col} ({op}): {val1}")
        else:
            filtre_aciklamalari.append("- Filtre uygulanmadı.")

        try:
            doc = Document()
            doc.add_heading("Filtreli Veri Çıktısı", 0)
            for aciklama in filtre_aciklamalari:
                doc.add_paragraph(aciklama)
            
            if not df_to_save.empty:
                table = doc.add_table(rows=1, cols=len(df_to_save.columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, column_name in enumerate(df_to_save.columns):
                    hdr_cells[i].text = str(column_name)
                for _, row in df_to_save.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
            else:
                doc.add_paragraph("Filtre sonrası veri bulunamadı.")
            
            doc.save(kayit_yolu)
            messagebox.showinfo("Kayıt", f"Veri '{yeni_ad}' olarak başarıyla kaydedildi.")
            self._dosya_agacini_guncelle()
        except Exception as e:
            messagebox.showerror("Hata", f"Kayıt başarısız: {e}")

    def veri_duzenle_pencere(self):
        # Bu fonksiyonun içeriği karmaşıklığı nedeniyle şimdilik boş bırakılmıştır.
        # İstenirse ayrı bir Toplevel penceresi içinde veri düzenleme arayüzü oluşturulabilir.
        messagebox.showinfo("Bilgi", "Veri düzenleme özelliği henüz aktif değil.")

    def _tabloyu_goster(self, df):
        if df is None:
            self.tree.delete(*self.tree.get_children())
            self.tree['columns'] = []
            return
            
        df_display = df.copy()
        for col in df_display.columns:
            if self._kolon_tipi_bul(col) == 'tarih':
                try:
                    df_display[col] = pd.to_datetime(df_display[col], dayfirst=True, errors='coerce').dt.strftime('%d.%m.%Y')
                except Exception: pass
        
        self.tree.delete(*self.tree.get_children())
        self.tree['columns'] = list(df_display.columns)
        for col in df_display.columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=120, anchor=tk.W)
        for _, row in df_display.iterrows():
            self.tree.insert('', tk.END, values=list(row))

if __name__ == "__main__":
    app = ExcelYoneticiV2()
    app.mainloop()

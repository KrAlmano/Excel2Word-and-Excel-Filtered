from optparse import Values
import os
import shutil
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import pandas as pd
import openpyxl
import re
import numpy as np
import sys
from docx import Document
import customtkinter as ctk


DATABASE_DIR = os.path.join(os.path.expanduser('~/Desktop'), 'ExcelVeritabaniV2')

class ExcelYoneticiV2(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Excel Yöneticisi V2")
        self.geometry("1300x800")
        ctk.set_appearance_mode("System")  # veya "Dark", "Light"
        ctk.set_default_color_theme("blue")
        self.df = None
        self.df_filtered = None
        self.secili_dosya = None
        self.secili_klasor = None
        self.filtreler = []  # [(kolon, operator, deger1, deger2, tip)]
        self._arayuz_olustur()
        self._dosya_agacini_guncelle()
        # Sağ tık menüsü dosya ve klasör için
        self.sag_menu = tk.Menu(self, tearoff=0)
        self.sag_menu.add_command(label="Sil", command=self._dosya_sil)
        self.sag_menu.add_command(label="Yeniden Adlandır", command=self._dosya_yeniden_adlandir)

    def _arayuz_olustur(self):
        ana_frame = ctk.CTkFrame(self)
        ana_frame.pack(fill=ctk.BOTH, expand=True)

        # Sol panel: Dosya/klasör ağacı
        sol_frame = ctk.CTkFrame(ana_frame)
        sol_frame.pack(side=ctk.LEFT, fill=ctk.Y, padx=10, pady=10)
        ctk.CTkLabel(sol_frame, text="Veritabanı", font=("Arial", 12, "bold")).pack(pady=(0,5))
        self.dosya_tree = ttk.Treeview(sol_frame, show="tree")  # Treeview için customtkinter alternatifi yok, ttk kullanılacak
        self.dosya_tree.pack(fill=ctk.Y, expand=True)
        self.dosya_tree.bind('<<TreeviewSelect>>', self._dosya_agacinda_secildi)
        self.dosya_tree.bind('<Button-3>', self._treeview_sag_tik_menu)  # Windows/Linux
        ctk.CTkButton(sol_frame, text="Excel Yükle", command=self.dosya_yukle).pack(pady=10)

        # Orta panel: Filtre ve tablo
        orta_frame = ctk.CTkFrame(ana_frame)
        orta_frame.pack(side=ctk.LEFT, fill=ctk.BOTH, expand=True, padx=10, pady=10)
        self._olustur_filtre_paneli(orta_frame)
        self.tablo_frame = ctk.CTkFrame(orta_frame)
        self.tablo_frame.pack(fill=ctk.BOTH, expand=True)
        self.tree = ttk.Treeview(self.tablo_frame, show="headings")  # Treeview için ttk kullanılacak
        self.tree.pack(side=ctk.LEFT, fill=ctk.BOTH, expand=True)
        self.vsb = ttk.Scrollbar(self.tablo_frame, orient="vertical", command=self.tree.yview)
        self.vsb.pack(side=ctk.RIGHT, fill=ctk.Y)
        self.hsb = ttk.Scrollbar(self.tablo_frame, orient="horizontal", command=self.tree.xview)
        self.hsb.pack(fill=ctk.X)
        self.tree.configure(yscrollcommand=self.vsb.set, xscrollcommand=self.hsb.set)
    


    def _dosya_agacini_guncelle(self):
        self.dosya_tree.delete(*self.dosya_tree.get_children())
        if not os.path.exists(DATABASE_DIR):
            os.makedirs(DATABASE_DIR)
        for klasor in sorted(os.listdir(DATABASE_DIR)):
            klasor_yol = os.path.join(DATABASE_DIR, klasor)
            if os.path.isdir(klasor_yol):
                kid = self.dosya_tree.insert('', 'end', text=klasor, open=True)
                for dosya in sorted(os.listdir(klasor_yol)):
                    if dosya.endswith('.xlsx'):
                        self.dosya_tree.insert(kid, 'end', text=dosya, values=(os.path.join(klasor, dosya),))

    def dosya_yukle(self):
        dosya = filedialog.askopenfilename(title="Excel Dosyası Seç", filetypes=[("Excel Dosyaları", "*.xlsx")])
        if not dosya:
            return
        dosya_adi = os.path.basename(dosya)
        ana_isim = os.path.splitext(dosya_adi)[0]
        hedef_klasor = os.path.join(DATABASE_DIR, ana_isim)
        if not os.path.exists(hedef_klasor):
            os.makedirs(hedef_klasor)
        hedef_yol = os.path.join(hedef_klasor, dosya_adi)
        try:
            shutil.copy2(dosya, hedef_yol)
            messagebox.showinfo("Başarılı", f"{dosya_adi} başarıyla yüklendi.")
            self._dosya_agacini_guncelle()
        except Exception as e:
            messagebox.showerror("Hata", f"Dosya yüklenemedi:\n{e}")

    def _dosya_agacinda_secildi(self, event):
        secim = self.dosya_tree.selection()
        if not secim:
            return
        item = self.dosya_tree.item(secim[0])
        parent = self.dosya_tree.parent(secim[0])
        if parent:  # Dosya seçildi
            klasor = self.dosya_tree.item(parent)['text']
            dosya_adi = item['text']
            dosya_yolu = os.path.join(DATABASE_DIR, klasor, dosya_adi)
            self.secili_klasor = klasor
            self.secili_dosya = dosya_adi
            try:
                df = pd.read_excel(dosya_yolu, engine="openpyxl")
                df = df.dropna(how='all')
                if any([str(col).startswith('Unnamed') for col in df.columns]):
                    new_header = df.iloc[0]
                    df = df[1:]
                    df.columns = new_header
                    df = df.reset_index(drop=True)
                self.df = df
                self.df_filtered = None
                self.filtreler = []
                self._tabloyu_goster(df)
                self._filtre_paneli_guncelle()
            except Exception as e:
                messagebox.showerror("Hata", f"Dosya okunamadı:\n{e}")
        else:
            self.secili_klasor = item['text']
            self.secili_dosya = None

    def _treeview_sag_tik_menu(self, event):
        iid = self.dosya_tree.identify_row(event.y)
        if iid:
            self.dosya_tree.selection_set(iid)
            try:
                self.sag_menu.tk_popup(event.x_root, event.y_root)
            finally:
                self.sag_menu.grab_release()
        else:
            return

    def _dosya_sil(self):
        secim = self.dosya_tree.selection()
        if not secim or len(secim) != 1:
            messagebox.showwarning("Uyarı!", "Lütfen silmek için bir dosya veya klasör seçin.")
            return
        item = self.dosya_tree.item(secim[0])
        parent = self.dosya_tree.parent(secim[0])
        if parent:  # Dosya sil
            klasor = self.dosya_tree.item(parent)['text']
            dosya_adi = item['text']
            dosya_yolu = os.path.join(DATABASE_DIR, klasor, dosya_adi)
            if messagebox.askyesno("Sil", f"{dosya_adi} dosyasını silmek istediğinize emin misiniz?"):
                try:
                    os.remove(dosya_yolu)
                    self._dosya_agacini_guncelle()
                except Exception as e:
                    messagebox.showerror("Hata", f"Dosya silinemedi:\n{e}")
        else:  # Klasör sil
            klasor = item['text']
            klasor_yolu = os.path.join(DATABASE_DIR, klasor)
            if messagebox.askyesno("Sil", f"{klasor} klasörünü ve içindeki tüm dosyaları silmek istediğinize emin misiniz?"):
                try:
                    shutil.rmtree(klasor_yolu)
                    self._dosya_agacini_guncelle()
                except Exception as e:
                    messagebox.showerror("Hata", f"Klasör silinemedi:\n{e}")

    def _dosya_yeniden_adlandir(self):
        secim = self.dosya_tree.selection()
        if not secim or len(secim) != 1:
            messagebox.showwarning("Uyarı", "Lütfen yeniden adlandırmak için bir dosya veya klasör seçin.")
            return
        item = self.dosya_tree.item(secim[0])
        parent = self.dosya_tree.parent(secim[0])
        if parent:  # Dosya yeniden adlandır
            klasor = self.dosya_tree.item(parent)['text']
            eski_ad = item['text']
            yeni_ad = simpledialog.askstring("Yeniden Adlandır", f"Yeni dosya adını girin (uzantı dahil):", initialvalue=eski_ad)
            if yeni_ad and yeni_ad != eski_ad:
                eski_yol = os.path.join(DATABASE_DIR, klasor, eski_ad)
                yeni_yol = os.path.join(DATABASE_DIR, klasor, yeni_ad)
                try:
                    os.rename(eski_yol, yeni_yol)
                    self._dosya_agacini_guncelle()
                except Exception as e:
                    messagebox.showerror("Hata", f"Dosya yeniden adlandırılamadı:\n{e}")
        else:  # Klasör yeniden adlandır
            eski_klasor = item['text']
            yeni_klasor = simpledialog.askstring("Yeniden Adlandır", f"Yeni klasör adını girin:", initialvalue=eski_klasor)
            if yeni_klasor and yeni_klasor != eski_klasor:
                eski_yol = os.path.join(DATABASE_DIR, eski_klasor)
                yeni_yol = os.path.join(DATABASE_DIR, yeni_klasor)
                try:
                    os.rename(eski_yol, yeni_yol)
                    self._dosya_agacini_guncelle()
                except Exception as e:
                    messagebox.showerror("Hata", f"Klasör yeniden adlandırılamadı:\n{e}")

    def _olustur_filtre_paneli(self, parent):
        self.filtre_frame = ctk.CTkFrame(parent)
        self.filtre_frame.pack(fill=ctk.X, pady=5)
        ctk.CTkLabel(self.filtre_frame, text="Filtreleme (Gelişmiş)").grid(row=0, column=0, columnspan=4, sticky="w", padx=5, pady=2)
        self.filtre_sutun = ctk.CTkComboBox(self.filtre_frame, state="disabled", width=180)
        self.filtre_sutun.grid(row=1, column=0, padx=5, pady=2)
        self.filtre_operator = ctk.CTkComboBox(self.filtre_frame, state="disabled", width=120)
        self.filtre_operator.grid(row=1, column=1, padx=5, pady=2)
        self.filtre_giris1 = ctk.CTkEntry(self.filtre_frame, width=120)
        self.filtre_giris1.grid(row=1, column=2, padx=5, pady=2)
        self.filtre_giris2 = ctk.CTkEntry(self.filtre_frame, width=120)
        self.filtre_giris2.grid(row=1, column=3, padx=5, pady=2)
        self.filtre_giris2.grid_remove()
        self.filtre_btn = ctk.CTkButton(self.filtre_frame, text="Filtre Ekle", command=self.filtre_ekle, state=ctk.DISABLED)
        self.filtre_btn.grid(row=2, column=0, padx=5, pady=2, sticky="ew")
        self.filtre_temizle_btn = ctk.CTkButton(self.filtre_frame, text="Tüm Filtreleri Temizle", command=self.filtre_temizle, state=ctk.DISABLED)
        self.filtre_temizle_btn.grid(row=2, column=1, padx=5, pady=2, sticky="ew")
        self.filtre_kaydet_btn = ctk.CTkButton(self.filtre_frame, text="Filtreli Veriyi Kaydet", command=self.filtreli_veriyi_kaydet, state=ctk.NORMAL)
        self.filtre_kaydet_btn.grid(row=2, column=2, padx=5, pady=2, sticky="ew")
        self.veri_duzenle_btn = ctk.CTkButton(self.filtre_frame, text="Veri Düzenle", command=self.veri_duzenle_pencere, state=ctk.NORMAL)
        self.veri_duzenle_btn.grid(row=2, column=3, padx=5, pady=2, sticky="ew")
        self.filtre_listbox = tk.Listbox(self.filtre_frame, width=80, height=3)  # Listbox için customtkinter alternatifi yok
        self.filtre_listbox.grid(row=3, column=0, columnspan=4, padx=5, pady=2, sticky="ew")
        self.filtre_sil_btn = ctk.CTkButton(self.filtre_frame, text="Seçili Filtreyi Kaldır", command=self.filtre_sil, state=ctk.DISABLED)
        self.filtre_sil_btn.grid(row=4, column=0, padx=5, pady=2, sticky="ew")
        self.filtre_sonuc_label = ctk.CTkLabel(self.filtre_frame, text="")
        self.filtre_sonuc_label.grid(row=4, column=1, columnspan=3, sticky="w")


    def _filtre_paneli_guncelle(self):
        if self.df is not None:
            self.filtre_sutun.configure(state="readonly")
            self.filtre_sutun.configure(values=list(self.df.columns))
            self.filtre_sutun.set("")
            self.filtre_operator.configure(state="readonly")
            self.filtre_operator.set("")
            self.filtre_giris1.delete(0, tk.END)
            self.filtre_giris2.delete(0, tk.END)
            self.filtre_btn.configure(state="normal")
            self.filtre_temizle_btn.configure(state="normal")
            self.filtre_sil_btn.configure(state="normal")
            self.filtre_sutun.bind('<<ComboboxSelected>>', self._filtre_sutun_degisti)
        else:
            self.filtre_sutun.configure(state="disabled")
            self.filtre_operator.configure(state="disabled")
            self.filtre_btn.configure(state="disabled")
            self.filtre_temizle_btn.configure(state="disabled")
            self.filtre_sil_btn.configure(state="disabled")

    def _filtre_sutun_degisti(self, event):
        col = self.filtre_sutun.get()
        if self.df is None or col not in self.df.columns:
            return
        # Sadece iki filtre türü: Aralık ve Anahtar Kelime
        self.filtre_operator.configure(values=['Aralık', 'Anahtar Kelime'])
        self.filtre_operator.set("")
        self.filtre_giris1.delete(0, tk.END)
        self.filtre_giris2.delete(0, tk.END)
        self.filtre_giris2.grid_remove()
        self.filtre_operator.bind('<<ComboboxSelected>>', self._filtre_operator_degisti)

    def _filtre_operator_degisti(self, event):
        op = self.filtre_operator.get()
        if op == 'Aralık':
            self.filtre_giris2.grid(row=0, column=3, padx=5, pady=2)
        else:
            self.filtre_giris2.grid_remove()

    def _kolon_tipi_bul(self, col):
        seri = self.df[col].dropna().astype(str)
        # Tarih mi?
        try:
            pd.to_datetime(seri.iloc[:10], dayfirst=True, errors='raise')
            return 'tarih'
        except:
            pass
        # Sayı mı?
        try:
            pd.to_numeric(seri.str.replace(r'[^0-9,.-]', '', regex=True).str.replace('.', '', regex=False).str.replace(',', '.', regex=False), errors='raise')
            return 'sayi'
        except:
            pass
        return 'metin'

    def filtre_ekle(self):
        col = self.filtre_sutun.get()
        op = self.filtre_operator.get()
        val1 = self.filtre_giris1.get()
        val2 = self.filtre_giris2.get() if op == 'Aralık' else None
        if not col or not op or not val1 or (op == 'Aralık' and not val2):
            messagebox.showwarning("Uyarı", "Lütfen tüm filtre alanlarını doldurun.")
            return
        tip = self._kolon_tipi_bul(col)
        self.filtreler.append((col, op, val1, val2, tip))
        self._filtre_listbox_guncelle()
        self._filtreleri_uygula()

    def _filtre_listbox_guncelle(self):
        self.filtre_listbox.delete(0, tk.END)
        for f in self.filtreler:
            col, op, val1, val2, tip = f
            if op == 'Aralık':
                self.filtre_listbox.insert(tk.END, f"{col} {op}: {val1} - {val2}")
            else:
                self.filtre_listbox.insert(tk.END, f"{col} {op} {val1}")

    def filtre_sil(self):
        sec = self.filtre_listbox.curselection()
        if not sec:
            return
        idx = sec[0]
        self.filtreler.pop(idx)
        self._filtre_listbox_guncelle()
        self._filtreleri_uygula()

    def filtre_temizle(self):
        self.filtreler = []
        self._filtre_listbox_guncelle()
        self._filtreleri_uygula()

    def _filtreleri_uygula(self):
        if self.df is None or not self.filtreler:
            self.df_filtered = None
            self._tabloyu_goster(self.df)
            self.filtre_sonuc_label.config(text="Toplam satır: {}".format(len(self.df)))
            return
        df = self.df.copy()
        for col, op, val1, val2, tip in self.filtreler:
            try:
                if op == 'Aralık':
                    if tip == 'tarih':
                        df[col] = pd.to_datetime(df[col], dayfirst=True, errors='coerce')
                        v1 = pd.to_datetime(val1, dayfirst=True, errors='coerce')
                        v2 = pd.to_datetime(val2, dayfirst=True, errors='coerce')
                        df = df[(df[col] >= v1) & (df[col] <= v2)]
                    else:
                        df[col] = df[col].astype(str).str.replace(r'[^0-9,.-]', '', regex=True)
                        df[col] = df[col].str.replace('.', '', regex=False)
                        df[col] = df[col].str.replace(',', '.', regex=False)
                        df[col] = pd.to_numeric(df[col], errors='coerce')
                        v1 = float(val1.replace('.', '').replace(',', '.'))
                        v2 = float(val2.replace('.', '').replace(',', '.'))
                        df = df[(df[col] >= v1) & (df[col] <= v2)]
                elif op == 'Anahtar Kelime':
                    keywords = re.split(r',| ve | ya da | veya ', val1)
                    keywords = [k.strip() for k in keywords if k.strip()]
                    mask = df[col].astype(str).apply(lambda x: all(kw.lower() in x.lower() for kw in keywords))
                    df = df[mask]
            except Exception as e:
                messagebox.showerror("Filtreleme Hatası", f"{col} filtresi uygulanamadı:\n{e}")
        self.df_filtered = df
        self._tabloyu_goster(df)
        self.filtre_sonuc_label.config(text="Filtreli satır: {} / Toplam: {}".format(len(df), len(self.df)))

    def filtreli_veriyi_kaydet(self):
        if self.df_filtered is None or self.secili_klasor is None or self.secili_dosya is None:
            messagebox.showwarning("Uyarı", "Kaydedilecek filtreli veri yok.")
            return
        ana_isim = self.secili_klasor
        dosya_adi = self.secili_dosya.replace('.xlsx', '_filtrede.docx')
        klasor_yolu = os.path.join(DATABASE_DIR, ana_isim)
        kayit_yolu = os.path.join(klasor_yolu, dosya_adi)
        sayac = 1
        while os.path.exists(kayit_yolu):
            dosya_adi = self.secili_dosya.replace('.xlsx', f'_filtrede_{sayac}.docx')
            kayit_yolu = os.path.join(klasor_yolu, dosya_adi)
            sayac += 1
        # Filtre açıklamalarını oluştur
        filtre_aciklamalari = []
        for f in self.filtreler:
            col, op, val1, val2, tip = f
            if op == 'Aralık':
                filtre_aciklamalari.append(f"Filtre: {col} aralığı: {val1} - {val2}")
            else:
                filtre_aciklamalari.append(f"Filtre: {col} anahtar kelime: {val1}")
        try:
            doc = Document()
            doc.add_heading("Filtreli Veri Çıktısı", 0)
            for aciklama in filtre_aciklamalari:
                doc.add_paragraph(aciklama)
            if self.df_filtered is not None and not self.df_filtered.empty:
                table = doc.add_table(rows=1, cols=len(self.df_filtered.columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, column_name in enumerate(self.df_filtered.columns):
                    hdr_cells[i].text = str(column_name)
                for _, row in self.df_filtered.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
            else:
                doc.add_paragraph("Filtre sonrası veri bulunamadı.")
            

            
            doc.save(kayit_yolu)
            messagebox.showinfo("Kayıt", f"Filtreli veri {dosya_adi} olarak Word'e kaydedildi.")
            self._dosya_agacini_guncelle()
        except Exception as e:
            messagebox.showerror("Hata", f"Kayıt başarısız: {e}")

    def veri_duzenle_pencere(self):
        if self.df is None and self.df_filtered is None:
            messagebox.showwarning("Uyarı", "Düzenlenecek veri yok.")
            return
        df = self.df_filtered if self.df_filtered is not None else self.df
        pencere = tk.Toplevel(self)
        pencere.title("Veri Düzenle")
        pencere.geometry("1000x600")
        # Filtre paneli
        filtre_frame = ctk.CTkFrame(pencere)
        filtre_frame.pack(fill=tk.X, pady=5)
        ctk.CTkLabel(filtre_frame, text="Filtreleme (Geçici)").pack(anchor="w", padx=5, pady=(0,2))
        filtre_sutun = ttk.Combobox(filtre_frame, state="readonly", width=20, values=list(df.columns))
        filtre_sutun.pack(side=tk.LEFT, padx=5)
        filtre_tur = ctk.CTkComboBox(filtre_frame, width=150)
        filtre_tur.configure(values=["Aralık", "Anahtar Kelime"])
        filtre_tur.pack(side=tk.LEFT, padx=5)
        filtre_giris1 = ttk.Entry(filtre_frame, width=15)
        filtre_giris1.pack(side=tk.LEFT, padx=5)
        filtre_giris2 = ttk.Entry(filtre_frame, width=15)
        filtre_giris2.pack(side=tk.LEFT, padx=5)
        filtre_giris2.pack_forget()
        def filtre_tur_degisti(evt):
            tur = filtre_tur.get()
            if tur == "Aralık":
                filtre_giris2.pack(side=tk.LEFT, padx=5)
                filtre_giris1.delete(0, tk.END)
                filtre_giris2.delete(0, tk.END)
                filtre_giris1.config(width=15)
                filtre_giris2.config(width=15)
            else:
                filtre_giris2.pack_forget()
                filtre_giris1.delete(0, tk.END)
                filtre_giris1.config(width=30)
        filtre_tur.bind('<<ComboboxSelected>>', filtre_tur_degisti)
        ttk.Button(filtre_frame, text="Filtrele", command=lambda: satirlari_guncelle()).pack(side=tk.LEFT, padx=5)
        # Satır listesi
        ttk.Label(pencere, text="Satır Seç:").pack(pady=5)
        satir_listbox = tk.Listbox(pencere, width=120, height=12)
        satir_listbox.pack(pady=5)
        gosterilen_indexler = list(df.index)
        def satirlari_guncelle():
            nonlocal gosterilen_indexler
            temp_df = df.copy()
            col = filtre_sutun.get()
            tur = filtre_tur.get()
            val1 = filtre_giris1.get()
            val2 = filtre_giris2.get() if tur == "Aralık" else None
            if col and tur and val1 and (tur != "Aralık" or val2):
                try:
                    if tur == "Aralık":
                        temp_df[col] = temp_df[col].astype(str).str.replace(r'[^0-9,.-]', '', regex=True)
                        temp_df[col] = temp_df[col].str.replace('.', '', regex=False)
                        temp_df[col] = temp_df[col].str.replace(',', '.', regex=False)
                        temp_df[col] = pd.to_numeric(temp_df[col], errors='coerce')
                        v1 = float(val1.replace('.', '').replace(',', '.'))
                        v2 = float(val2.replace('.', '').replace(',', '.'))
                        temp_df = temp_df[(temp_df[col] >= v1) & (temp_df[col] <= v2)]
                    else:
                        keywords = re.split(r',| ve | ya da | veya ', val1)
                        keywords = [k.strip() for k in keywords if k.strip()]
                        mask = temp_df[col].astype(str).apply(lambda x: all(kw.lower() in x.lower() for kw in keywords))
                        temp_df = temp_df[mask]
                except Exception:
                    pass
            satir_listbox.delete(0, tk.END)
            gosterilen_indexler = list(temp_df.index)
            for i, row in temp_df.iterrows():
                satir_listbox.insert(tk.END, f"{i}: {list(row)}")
        satirlari_guncelle()
        # Entryler
        duzen_frame = ttk.Frame(pencere)
        duzen_frame.pack(pady=10)
        entryler = []
        for idx, col in enumerate(df.columns):
            ttk.Label(duzen_frame, text=col).grid(row=0, column=idx)
            e = ttk.Entry(duzen_frame, width=15)
            e.grid(row=1, column=idx)
            entryler.append(e)
        def satir_secildi(evt):
            sec = satir_listbox.curselection()
            if not sec:
                return
            row_idx = sec[0]
            orijinal_index = gosterilen_indexler[row_idx]
            for i, col in enumerate(df.columns):
                entryler[i].delete(0, tk.END)
                entryler[i].insert(0, str(df.loc[orijinal_index, col]))
        satir_listbox.bind('<<ListboxSelect>>', satir_secildi)
        def ekle():
            yeni = [e.get() for e in entryler]
            if any(yeni):
                yeni_df = pd.DataFrame([yeni], columns=df.columns)
                self.df = pd.concat([self.df, yeni_df], ignore_index=True)
                if self.df_filtered is not None:
                    self.df_filtered = pd.concat([self.df_filtered, yeni_df], ignore_index=True)
                self._tabloyu_goster(self.df_filtered if self.df_filtered is not None else self.df)
                pencere.destroy()
        def sil():
            sec = satir_listbox.curselection()
            if not sec:
                return
            row_idx = sec[0]
            orijinal_index = gosterilen_indexler[row_idx]
            self.df = self.df.drop(index=orijinal_index).reset_index(drop=True)
            if self.df_filtered is not None:
                self.df_filtered = self.df_filtered.drop(index=orijinal_index).reset_index(drop=True)
            self._tabloyu_goster(self.df_filtered if self.df_filtered is not None else self.df)
            pencere.destroy()
        def guncelle():
            sec = satir_listbox.curselection()
            if not sec:
                return
            row_idx = sec[0]
            orijinal_index = gosterilen_indexler[row_idx]
            for i, col in enumerate(df.columns):
                self.df.at[orijinal_index, col] = entryler[i].get()
                if self.df_filtered is not None:
                    self.df_filtered.at[orijinal_index, col] = entryler[i].get()
            self._tabloyu_goster(self.df_filtered if self.df_filtered is not None else self.df)
            pencere.destroy()
        btn_frame = ttk.Frame(pencere)
        btn_frame.pack(pady=10)
        ttk.Button(btn_frame, text="Ekle", command=ekle).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Sil", command=sil).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Güncelle", command=guncelle).pack(side=tk.LEFT, padx=10)

    def _tabloyu_goster(self, df):
        # Tarih sütunlarını düzgün formatla göster
        for col in df.columns:
            if self.df is not None and self._kolon_tipi_bul(col) == 'tarih':
                try:
                    df[col] = pd.to_datetime(df[col], dayfirst=True, errors='coerce').dt.strftime('%d.%m.%Y')
                except Exception:
                    pass
        self.tree.delete(*self.tree.get_children())
        self.tree['columns'] = list(df.columns)
        for col in df.columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=120, anchor=tk.W)
        for _, row in df.iterrows():
            self.tree.insert('', tk.END, values=list(row))
        self.update_idletasks()
        for col in df.columns:
            self.tree.column(col, width=max(120, min(300, self.tree.winfo_width() // len(df.columns))))

if __name__ == "__main__":
    app = ExcelYoneticiV2()
    app.mainloop()